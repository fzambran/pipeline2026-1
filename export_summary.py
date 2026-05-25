"""
Genera el resumen del pipeline en formato Word (.docx) y texto plano (.txt).
Salida: data/reports/pipeline_summary.docx y data/reports/pipeline_summary.txt
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = Path("data/reports")

STAGES = [
    (
        "Stage 1 — Limpieza (src/clean.py)",
        (
            "Elimina las 19 columnas 'zero*' de relleno, decodifica Sex (0/1 → male/female) "
            "y Embarked (0/1/2 → S/C/Q), imputa nulos, filtra rangos inválidos y crea 4 columnas "
            "nuevas (family_size, is_alone, fare_per_person, age_group). "
            "Resultado: 1 309 → 1 309 filas, 28 → 13 columnas."
        ),
    ),
    (
        "Stage 2 — Validación semántica (src/validate.py)",
        (
            "Aplica 12 reglas (dominio, rango, unicidad, semántica). Detectó 2 niños "
            "(5 y 11 años) marcados viajando solos como anomalías semánticas. "
            "Resultado: 1 307 válidos, 2 rechazados. "
            "Reporte en data/reports/validation_report.txt."
        ),
    ),
    (
        "Stage 3 — Carga a PostgreSQL (src/load.py)",
        (
            "Crea la tabla passengers con 8 constraints en PostgreSQL, inserta fila a fila "
            "manejando UniqueViolation y CheckViolation, separa insertados de rechazados, "
            "y genera logs/load.log con el detalle de cada operación."
        ),
    ),
]

FILES_TABLE = [
    ("src/clean.py",           "Stage 1: limpieza"),
    ("src/validate.py",        "Stage 2: validación semántica"),
    ("src/load.py",            "Stage 3: carga a PostgreSQL"),
    ("main.py",                "Orquestador con --stage opcional"),
    ("sql/create_table.sql",   "DDL de la tabla passengers con 8 constraints"),
    ("docker-compose.yml",     "postgres 16 + pgAdmin 4 + contenedor pipeline"),
    ("Dockerfile",             "Imagen del pipeline (python:3.13-slim + uv)"),
    (".env",                   "Variables de conexión para desarrollo local"),
    ("README.md",              "Documentación completa"),
]

COMMANDS = [
    "# Infraestructura + pipeline completo en contenedores",
    "docker compose up --build",
    "",
    "# O localmente (con postgres en Docker)",
    "docker compose up postgres pgadmin -d",
    "uv run python main.py",
    "",
    "# Ejecutar una sola etapa",
    "uv run python main.py --stage clean",
    "uv run python main.py --stage validate",
    "uv run python main.py --stage load",
]


# ──────────────────────────────────────────────
# Word
# ──────────────────────────────────────────────

def _heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def _add_table(doc: Document, headers: list[str], rows: list[tuple]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build_word(out_path: Path) -> None:
    doc = Document()

    # Title
    title = doc.add_heading("Resumen del Pipeline — Titanic Data Pipeline", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Archivos creados
    _heading(doc, "Archivos creados", level=1)
    _add_table(doc, ["Archivo", "Rol"], FILES_TABLE)
    doc.add_paragraph()

    # Etapas
    _heading(doc, "Descripción de etapas", level=1)
    for stage_title, stage_body in STAGES:
        _heading(doc, stage_title, level=2)
        doc.add_paragraph(stage_body)

    # Comandos de ejecución
    _heading(doc, "Comandos de ejecución", level=1)
    for line in COMMANDS:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        if line.startswith("#"):
            run.font.color.rgb = RGBColor(0x6A, 0x6A, 0x6A)

    # pgAdmin
    _heading(doc, "pgAdmin", level=1)
    doc.add_paragraph(
        "Disponible en http://localhost:8080 con las credenciales definidas en el archivo .env. "
        "Al agregar el servidor usar 'postgres' como host (nombre del servicio en la red Docker)."
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"Word guardado en: {out_path}")


# ──────────────────────────────────────────────
# Texto plano
# ──────────────────────────────────────────────

def build_txt(out_path: Path) -> None:
    sep = "=" * 64
    lines: list[str] = []

    lines += [
        sep,
        "  RESUMEN DEL PIPELINE — TITANIC DATA PIPELINE",
        sep,
        "",
    ]

    # Archivos creados
    lines += ["ARCHIVOS CREADOS", "-" * 40]
    col_w = max(len(f) for f, _ in FILES_TABLE) + 2
    for fname, rol in FILES_TABLE:
        lines.append(f"  {fname:<{col_w}}{rol}")
    lines.append("")

    # Etapas
    lines += ["DESCRIPCIÓN DE ETAPAS", "-" * 40]
    for stage_title, stage_body in STAGES:
        lines.append(f"\n  {stage_title}")
        # Wrap body at ~70 chars
        words = stage_body.split()
        current_line = "    "
        for word in words:
            if len(current_line) + len(word) + 1 > 72:
                lines.append(current_line.rstrip())
                current_line = "    " + word + " "
            else:
                current_line += word + " "
        if current_line.strip():
            lines.append(current_line.rstrip())
    lines.append("")

    # Comandos
    lines += ["COMANDOS DE EJECUCIÓN", "-" * 40]
    for cmd in COMMANDS:
        lines.append(f"  {cmd}" if cmd else "")
    lines.append("")

    # pgAdmin
    lines += [
        "PGADMIN",
        "-" * 40,
        "  Disponible en http://localhost:8080 con las credenciales del .env.",
        "  Usar 'postgres' como host al agregar el servidor.",
        "",
        sep,
    ]

    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text("\n".join(lines), encoding="utf-8")
    print(f"Texto plano guardado en: {out_path}")


if __name__ == "__main__":
    build_word(OUT_DIR / "pipeline_summary.docx")
    build_txt(OUT_DIR / "pipeline_summary.txt")
